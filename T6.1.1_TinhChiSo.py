#T6.1.1: Tính các chỉ số: Lợi nhuận ròng, Tỷ lệ lợi nhuận ròng, Tỷ lệ nợ, Tỷ lệ vốn chủ sở hữu:
import pandas as pd
from docx import Document
from docx.shared import Inches
# Đọc dữ liệu từ file Excel
df = pd.read_excel('data.xlsx')
# Tính toán các chỉ số tài chính
df['Lợi nhuận ròng'] = df['Doanh thu'] - df['Chi phí hàng bán'] - df['Chi phí bán hàng'] - df['Chi phí quản lý'] - df['Lãi vay'] - df['Thuế thu nhập doanh nghiệp']
df['Tỷ lệ lợi nhuận ròng'] = df['Lợi nhuận ròng'] / df['Doanh thu']
df['Tỷ lệ nợ'] = df['Nợ phải trả'] / df['Tổng tài sản']
df['Tỷ lệ vốn chủ sở hữu'] = df['Vốn chủ sở hữu'] / df['Tổng tài sản']
# Tạo một file Word mới
document = Document()
# Thêm tiêu đề vào tài liệu
document.add_heading('Báo cáo các chỉ số tài chính', 0)
# Thêm bảng kết quả vào tài liệu
table = document.add_table(rows=len(df) + 1, cols=len(df.columns))
# Đặt tiêu đề cho các cột
for col_idx, col_name in enumerate(df.columns):
    table.cell(0, col_idx).text = col_name
# Điền dữ liệu vào các ô
for i, row in df.iterrows():
    for j, cell in enumerate(row):
        table.cell(i+1, j).text = str(cell)
# Lưu file Word
document.save('KetQua.docx')
print("Đã lưu kết quả vào file KetQua.docx")
